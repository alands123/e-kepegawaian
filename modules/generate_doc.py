"""
generate_doc.py — Document Generation System (KP4 & KGB DOCX)
Uses python-docx to create professional government documents.
"""

import io
from datetime import datetime
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from modules.utils import format_date, format_currency, format_nip


def _set_cell_border(cell, **kwargs):
    """Set cell borders (utility)."""
    from docx.oxml.ns import qn
    from lxml import etree
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = etree.SubElement(tcPr, qn('w:tcBorders'))
    for edge, val in kwargs.items():
        element = etree.SubElement(tcBorders, qn(f'w:{edge}'))
        element.set(qn('w:val'), val.get('val', 'single'))
        element.set(qn('w:sz'), val.get('sz', '4'))
        element.set(qn('w:color'), val.get('color', '000000'))
        element.set(qn('w:space'), val.get('space', '0'))


def _add_styled_paragraph(doc, text, font_name="Times New Roman", size=12,
                          bold=False, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                          space_before=0, space_after=0, color=None):
    """Add a styled paragraph to the document."""
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    return p


def generate_kp4_doc(data: dict) -> io.BytesIO:
    """Generate a KP4 document as DOCX and return a BytesIO buffer."""
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2)

    # Header
    _add_styled_paragraph(doc, "PEMERINTAH DAERAH PROVINSI ...", size=12, bold=True,
                          alignment=WD_ALIGN_PARAGRAPH.CENTER)
    _add_styled_paragraph(doc, "DINAS PENDIDIKAN", size=12, bold=True,
                          alignment=WD_ALIGN_PARAGRAPH.CENTER)
    _add_styled_paragraph(doc, "Jl. ........ No. ........ Telp. ........", size=10,
                          alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

    # Line separator
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("_" * 75)
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0, 0, 0)

    # Title
    _add_styled_paragraph(doc, "", space_after=6)
    _add_styled_paragraph(doc, "KARTU PERMOHONAN PENAMBAHAN PENGHASILAN PEGAWAI", size=13, bold=True,
                          alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=12, space_after=4)
    _add_styled_paragraph(doc, f"Nomor: {data.get('nomor_surat', '...')}", size=11,
                          alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

    # Data table
    table = doc.add_table(rows=8, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    fields = [
        ("Nama", data.get("nama", "-")),
        ("NIP", format_nip(data.get("nip", "-"))),
        ("Golongan / Pangkat", f"{data.get('golongan', '-')} / {data.get('pangkat', '-')}"),
        ("Jabatan", data.get("jabatan", "-")),
        ("Unit Kerja", data.get("unit_kerja", "-")),
        ("TMT Golongan", format_date(data.get("tmt_golongan", ""))),
        ("Perihal", data.get("perihal", "Permohonan Penambahan Penghasilan Pegawai")),
        ("Keterangan", data.get("keterangan", "-")),
    ]

    for i, (label, value) in enumerate(fields):
        cell_label = table.cell(i, 0)
        cell_value = table.cell(i, 1)
        cell_label.text = label
        cell_value.text = str(value)

        for cell in [cell_label, cell_value]:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(11)

        # Bold labels
        for run in cell_label.paragraphs[0].runs:
            run.font.bold = True

    # Set column widths
    for row in table.rows:
        row.cells[0].width = Cm(5)
        row.cells[1].width = Cm(10)

    # Spacing
    _add_styled_paragraph(doc, "", space_after=24)

    # Signature area
    today = datetime.now()
    _add_styled_paragraph(doc, f"........, {format_date(today.strftime('%Y-%m-%d'))}", size=11,
                          alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    _add_styled_paragraph(doc, "Mengetahui,", size=11, alignment=WD_ALIGN_PARAGRAPH.RIGHT, space_before=4)
    _add_styled_paragraph(doc, "Kepala Dinas", size=11, alignment=WD_ALIGN_PARAGRAPH.RIGHT, space_before=4)
    _add_styled_paragraph(doc, "", space_after=48, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    _add_styled_paragraph(doc, "(........................)", size=11, bold=True,
                          alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    _add_styled_paragraph(doc, "NIP. ........................", size=11,
                          alignment=WD_ALIGN_PARAGRAPH.RIGHT)

    # Watermark text
    _add_styled_paragraph(doc, "", space_after=24)
    _add_styled_paragraph(doc, "[ DRAFT ]", size=9, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                          color=(180, 180, 180))

    # Save to buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def generate_kgb_doc(data: dict) -> io.BytesIO:
    """Generate a KGB document as DOCX and return a BytesIO buffer."""
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2)

    # Header
    _add_styled_paragraph(doc, "PEMERINTAH DAERAH PROVINSI ...", size=12, bold=True,
                          alignment=WD_ALIGN_PARAGRAPH.CENTER)
    _add_styled_paragraph(doc, "DINAS PENDIDIKAN", size=12, bold=True,
                          alignment=WD_ALIGN_PARAGRAPH.CENTER)
    _add_styled_paragraph(doc, "Jl. ........ No. ........ Telp. ........", size=10,
                          alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

    p = doc.add_paragraph()
    run = p.add_run("_" * 75)
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0, 0, 0)

    # Title
    _add_styled_paragraph(doc, "", space_after=6)
    _add_styled_paragraph(doc, "SURAT KENAIKAN GAJI BERKALA", size=14, bold=True,
                          alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=12, space_after=4)
    _add_styled_paragraph(doc, f"Nomor: {data.get('nomor_surat', '...')}", size=11,
                          alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=16)

    # Body
    _add_styled_paragraph(doc, "Berdasarkan peraturan perundang-undangan yang berlaku, dengan ini diberikan "
                          "Kenaikan Gaji Berkala kepada:", size=11, space_after=12)

    # Data table
    table = doc.add_table(rows=10, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    fields = [
        ("Nama", data.get("nama", "-")),
        ("NIP", format_nip(data.get("nip", "-"))),
        ("Golongan / Pangkat", f"{data.get('golongan', '-')} / {data.get('pangkat', '-')}"),
        ("Jabatan", data.get("jabatan", "-")),
        ("Unit Kerja", data.get("unit_kerja", "-")),
        ("TMT Golongan", format_date(data.get("tmt_golongan", ""))),
        ("Jatuh Tempo KGB", format_date(data.get("jatuh_tempo", ""))),
        ("Gaji Pokok Lama", format_currency(data.get("gaji_lama", "0"))),
        ("Gaji Pokok Baru", format_currency(data.get("gaji_baru", "0"))),
        ("Keterangan", data.get("keterangan", "-")),
    ]

    for i, (label, value) in enumerate(fields):
        cell_label = table.cell(i, 0)
        cell_value = table.cell(i, 1)
        cell_label.text = label
        cell_value.text = str(value)
        for cell in [cell_label, cell_value]:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(11)
        for run in cell_label.paragraphs[0].runs:
            run.font.bold = True

    for row in table.rows:
        row.cells[0].width = Cm(5)
        row.cells[1].width = Cm(10)

    # Closing
    _add_styled_paragraph(doc, "", space_after=12)
    _add_styled_paragraph(doc,
        "Demikian surat kenaikan gaji berkala ini dibuat untuk dapat dipergunakan sebagaimana mestinya.",
        size=11, space_after=24)

    # Signature
    today = datetime.now()
    _add_styled_paragraph(doc, f"........, {format_date(today.strftime('%Y-%m-%d'))}", size=11,
                          alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    _add_styled_paragraph(doc, "Kepala Dinas", size=11, alignment=WD_ALIGN_PARAGRAPH.RIGHT, space_before=4)
    _add_styled_paragraph(doc, "", space_after=48, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    _add_styled_paragraph(doc, "(........................)", size=11, bold=True,
                          alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    _add_styled_paragraph(doc, "NIP. ........................", size=11,
                          alignment=WD_ALIGN_PARAGRAPH.RIGHT)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
